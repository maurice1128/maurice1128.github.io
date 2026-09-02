# -*- coding: utf-8 -*-
"""r599: render the manuscript and supplement as .docx for submission.

BMC journals accept DOC, DOCX, RTF, TeX or PDF; they do not accept Markdown. Pandoc is not installed
on this machine, so this converts directly with python-docx. It handles the constructs this
manuscript actually uses: headings at three levels, paragraphs, bold and italic runs, inline code,
pipe tables, and horizontal rules. Everything is set in the layout BMC asks for: A4, 2.5 cm margins,
Times New Roman 12 pt, double spacing, and continuous line numbering, which reviewers use to point at
text.

Figures are not embedded. BMC wants them uploaded as separate files, with the legends in the
manuscript, which is how the source already has them.
"""
import io, os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PAP = r"C:\Users\maurice\Desktop\spasticity_paper\paper"
JOBS = [("MANUSCRIPT_r541.md", "MANUSCRIPT_r541.docx"),
        ("SUPPLEMENT_r541.md", "SUPPLEMENT_r541.docx")]

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`|<sup>.+?</sup>|<sub>.+?</sub>)", re.S)


def add_runs(par, text):
    """Split a paragraph into bold, italic, code and plain runs."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) > 4:
            r = par.add_run(piece[2:-2]); r.bold = True
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            r = par.add_run(piece[1:-1]); r.italic = True
        elif piece.startswith("<sup>") and piece.endswith("</sup>"):
            r = par.add_run(piece[5:-6]); r.font.superscript = True
        elif piece.startswith("<sub>") and piece.endswith("</sub>"):
            r = par.add_run(piece[5:-6]); r.font.subscript = True
        elif piece.startswith("`") and piece.endswith("`") and len(piece) > 2:
            r = par.add_run(piece[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10.5)
        else:
            par.add_run(piece)


def line_numbers(section, start=1):
    sectPr = section._sectPr
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:start"), str(start))
    ln.set(qn("w:restart"), "continuous")
    sectPr.append(ln)


def base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    for s in doc.sections:
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(2.5)
        line_numbers(s)


def render(src, dst):
    text = io.open(os.path.join(PAP, src), encoding="utf-8").read()
    doc = Document()
    base_style(doc)
    blocks = text.split("\n\n")
    i = 0
    n_tab = 0
    while i < len(blocks):
        blk = blocks[i].strip("\n")
        i += 1
        if not blk.strip():
            continue
        first = blk.split("\n")[0]

        if first.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(" ".join(first[2:].split()))
            r.bold = True
            r.font.size = Pt(16)
            p.paragraph_format.space_after = Pt(12)
            continue
        if first.startswith("#### "):
            p = doc.add_paragraph()
            add_runs(p, " ".join(first[5:].split()))
            for r in p.runs:
                r.bold = True
                r.italic = True
            p.paragraph_format.space_before = Pt(10)
            continue
        if first.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(" ".join(first[4:].split()))
            r.bold = True
            r.font.size = Pt(12.5)
            p.paragraph_format.space_before = Pt(12)
            continue
        if first.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(" ".join(first[3:].split()))
            r.bold = True
            r.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(16)
            continue
        if blk.strip() == "---":
            continue

        if first.lstrip().startswith("|"):
            rows = [x for x in blk.split("\n") if x.strip().startswith("|")]
            rows = [r for r in rows if not re.match(r"^\s*\|[\s:|-]+\|\s*$", r)]
            cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
            if not cells:
                continue
            width = max(len(c) for c in cells)
            t = doc.add_table(rows=len(cells), cols=width)
            t.style = "Table Grid"
            n_tab += 1
            for ri, row in enumerate(cells):
                for ci in range(width):
                    cell = t.cell(ri, ci)
                    cell.text = ""
                    par = cell.paragraphs[0]
                    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    add_runs(par, row[ci] if ci < len(row) else "")
                    for r in par.runs:
                        r.font.size = Pt(10)
                        if ri == 0:
                            r.bold = True
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        add_runs(p, " ".join(blk.split()))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    out = os.path.join(PAP, dst)
    doc.save(out)
    return out, n_tab


for src, dst in JOBS:
    path, tabs = render(src, dst)
    print("%-26s -> %-26s %6d bytes, %d tables"
          % (src, dst, os.path.getsize(path), tabs))
