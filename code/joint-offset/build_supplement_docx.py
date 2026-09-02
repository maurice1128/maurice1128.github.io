"""Build the supplementary-material .docx from SUPPLEMENT_SUBMISSION.md.

Input  : SUPPLEMENT_SUBMISSION.md
Output : ROWV_JBiomech_supplement.docx

Landscape, because 35 of the tables are verbatim result-file dumps and the widest runs to 178
characters. Portrait at any legible size would rewrap them, and the point of reproducing them
verbatim is that no number has been retyped -- a rewrapped column is a retyped column. Prose
carries a right indent so its measure stays readable on a landscape page; the dumps use the
full width.
"""
import io
import os
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SRC = "D:/ROWV_paper/SUPPLEMENT_SUBMISSION.md"
OUT = "D:/ROWV_paper/ROWV_JBiomech_supplement.docx"

doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Inches(11.69), Inches(8.27)      # A4 landscape
for m in ("left_margin", "right_margin"):
    setattr(sec, m, Inches(0.8))
for m in ("top_margin", "bottom_margin"):
    setattr(sec, m, Inches(0.7))

USABLE_IN = 11.69 - 1.6            # 10.09"
PROSE_INDENT = 3.0                 # keeps prose to a ~7" measure on a landscape page

st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(10.5)
st.paragraph_format.line_spacing = 1.15
st.paragraph_format.space_after = Pt(6)


def restyle(name, size, before=12, after=6):
    """Word ships the heading styles blue and sans-serif; match the document instead."""
    s = doc.styles[name]
    s.font.name = "Times New Roman"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.italic = False
    s.font.color.rgb = RGBColor(0, 0, 0)
    rpr = s.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        if rf.get(qn("w:" + a)) is not None:
            del rf.attrib[qn("w:" + a)]
    for a in ("ascii", "hAnsi", "cs"):
        rf.set(qn("w:" + a), "Times New Roman")
    pf = s.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.keep_with_next = True


restyle("Title", 16, before=0, after=12)
restyle("Heading 1", 13)
restyle("Heading 2", 11.5)
doc.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def runs(par, text):
    """Render **bold**, *italic*, `code`; strip markdown escapes."""
    text = text.replace("\\*", "\u0001")
    for piece in INLINE.split(text):
        if not piece:
            continue
        piece = piece.replace("\u0001", "*")
        if piece.startswith("**") and piece.endswith("**"):
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            par.add_run(piece[1:-1]).italic = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        else:
            par.add_run(piece)


def para(text):
    p = doc.add_paragraph()
    runs(p, text)
    p.paragraph_format.right_indent = Inches(PROSE_INDENT)
    return p


def heading(text, level):
    return doc.add_paragraph(text, style="Heading 1" if level == 1 else "Heading 2")


def code_block(lines):
    """A verbatim result-file dump: monospace, single-spaced, never rewrapped."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(8)
    pf.space_before = Pt(4)
    pf.left_indent = Inches(0.12)
    for i, ln in enumerate(lines):
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(6.4)
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.append(rf)
        rf.set(qn("w:ascii"), "Consolas")
        rf.set(qn("w:hAnsi"), "Consolas")
        if i < len(lines) - 1:
            r.add_break()
    return p


CHW = 0.052                        # inches per character, Times New Roman 8 pt
PAD = 0.14
MAXC = 60
_BREAKS = re.compile("[ /" + chr(0x2013) + chr(0x2014) + chr(0x2192) + "-]+")


def longest_token(s):
    return max([len(x) for x in _BREAKS.split(s) if x] or [1])


def col_widths(rows, ncol):
    """Every column at least as wide as its longest word, then share out the rest."""
    need = [max(longest_token(rows[i][j]) for i in range(len(rows))) * CHW * 1.08 + PAD
            for j in range(ncol)]
    want = [min(max(len(rows[i][j]) for i in range(len(rows))), MAXC) * CHW + PAD
            for j in range(ncol)]
    want = [max(w, n) for w, n in zip(want, need)]
    if sum(need) >= USABLE_IN:
        k = USABLE_IN / sum(need)
        return [Inches(n * k) for n in need]
    if sum(want) <= USABLE_IN:
        return [Inches(w) for w in want]
    slack = USABLE_IN - sum(need)
    room = sum(w - n for w, n in zip(want, need)) or 1.0
    return [Inches(n + slack * (w - n) / room) for n, w in zip(need, want)]


def _border(tag, on):
    e = OxmlElement("w:" + tag)
    e.set(qn("w:val"), "single" if on else "none")
    e.set(qn("w:sz"), "8" if on else "0")
    e.set(qn("w:color"), "000000")
    e.set(qn("w:space"), "0")
    return e


def md_table(rows):
    """Horizontal rules only: the Guide for Authors asks authors to avoid vertical rules."""
    ncol = max(len(r) for r in rows)
    rows = [r + [""] * (ncol - len(r)) for r in rows]
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = doc.styles["Normal Table"]
    t.autofit = False
    b = OxmlElement("w:tblBorders")
    for tag, on in (("top", True), ("bottom", True), ("left", False),
                    ("right", False), ("insideH", False), ("insideV", False)):
        b.append(_border(tag, on))
    t._tbl.tblPr.append(b)
    m = OxmlElement("w:tblCellMar")
    for tag, w in (("top", "30"), ("bottom", "30"), ("left", "80"), ("right", "80")):
        e = OxmlElement("w:" + tag)
        e.set(qn("w:w"), w)
        e.set(qn("w:type"), "dxa")
        m.append(e)
    t._tbl.tblPr.append(m)
    widths = col_widths(rows, ncol)
    for i, row in enumerate(rows):
        trPr = t.rows[i]._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))
        if i == 0:
            trPr.append(OxmlElement("w:tblHeader"))
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            c.width = widths[j]
            c.text = ""
            if i == 0:
                tb = OxmlElement("w:tcBorders")
                tb.append(_border("bottom", True))
                c._tc.get_or_add_tcPr().append(tb)
            p = c.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            runs(p, cell)
            for r in p.runs:
                r.font.size = Pt(8)
                if i == 0:
                    r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


src = io.open(SRC, encoding="utf-8").read()
lines = src.split("\n")
i = 0
buf, tbl, code = [], [], None


def flush_par():
    global buf
    if buf:
        para(" ".join(buf).strip())
        buf = []


def flush_tbl():
    global tbl
    if tbl:
        md_table(tbl)
        tbl = []


while i < len(lines):
    ln = lines[i].rstrip()
    if ln.startswith("```"):
        if code is None:
            flush_par()
            flush_tbl()
            code = []
        else:
            code_block(code)
            code = None
        i += 1
        continue
    if code is not None:
        code.append(lines[i].rstrip("\n"))
        i += 1
        continue
    if ln.startswith("|"):
        flush_par()
        cells = [c.strip() for c in ln.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
            tbl.append(cells)
        i += 1
        continue
    flush_tbl()
    if not ln.strip():
        flush_par()
        i += 1
        continue
    if ln.startswith("---"):
        flush_par()
        i += 1
        continue
    if ln.startswith("#"):
        flush_par()
        lvl = len(ln) - len(ln.lstrip("#"))
        txt = ln.lstrip("# ").strip()
        if lvl == 1:
            doc.add_paragraph(txt, style="Title")
        else:
            # every numbered supplementary table starts on a fresh page
            if re.match(r"(Table )?S\d+[a-z]?\.", txt):
                doc.add_page_break()
            heading(txt, 1 if lvl == 2 else 2)
        i += 1
        continue
    buf.append(ln.strip())
    i += 1

flush_par()
flush_tbl()
if code:
    code_block(code)

doc.save(OUT)
print("wrote", OUT, os.path.getsize(OUT), "bytes")
