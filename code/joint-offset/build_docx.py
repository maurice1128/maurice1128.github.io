"""Build the Journal of Biomechanics submission .docx from the markdown sources.

Inputs : MANUSCRIPT_SUBMISSION.md, TABLES_SUBMISSION.md, figures/*.png
Output : ROWV_JBiomech_submission.docx, submission_figures/Figure_N.png

Layout follows the J Biomech Guide for Authors: text, references, figure captions,
then tables, each table on its own page. The figures are NOT embedded. The guide
requires that artwork "must be supplied as separate files along with the manuscript",
each "as a separate file using a logical naming convention (for example, Figure_1,
Figure_2 etc)", so this script writes them out under that naming instead.

Headings use the built-in Heading styles rather than hand-bolded Normal paragraphs,
so Word's navigation pane and the publisher's structure check both see the sections.
"""
import re, io, os, sys
sys.path.insert(0, "D:/ROWV_paper")
from paper_stats import references as _references
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "D:/ROWV_paper/MANUSCRIPT_SUBMISSION.md"
TBL = "D:/ROWV_paper/TABLES_SUBMISSION.md"
FIGDIR = "D:/ROWV_paper/figures"
OUT = "D:/ROWV_paper/ROWV_JBiomech_submission.docx"
FIGOUT = "D:/ROWV_paper/submission_figures"

FIGS = [("Figure 1", "fig1_dissociation.png"),
        ("Figure 2", "fig2_perjoint.png"),
        ("Figure 3", "fig3_dose_distortion.png"),
        ("Figure 4", "fig5_posture.png")]

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)   # A4
for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(sec, m, Inches(1.0))

st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)
st.paragraph_format.line_spacing = 2.0          # J Biomech wants double spacing
st.paragraph_format.space_after = Pt(0)

def add_line_numbers(section):
    sectPr = section._sectPr
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1"); ln.set(qn("w:restart"), "continuous")
    sectPr.append(ln)
add_line_numbers(sec)

def restyle(name, size, before=12, after=6):
    """Make a built-in style match the manuscript: Word ships these blue and sans-serif."""
    st = doc.styles[name]
    st.font.name = "Times New Roman"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.italic = False
    st.font.color.rgb = RGBColor(0, 0, 0)
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for a in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        if rf.get(qn("w:" + a)) is not None:
            del rf.attrib[qn("w:" + a)]
    for a in ("ascii", "hAnsi", "cs"):
        rf.set(qn("w:" + a), "Times New Roman")
    pf = st.paragraph_format
    pf.line_spacing = 2.0
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.keep_with_next = True

restyle("Title", 16, before=0, after=12)
restyle("Heading 1", 13)
restyle("Heading 2", 12)
doc.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

_seen_table = []               # Table 1 shares the page with the 'Tables' heading

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")

def runs(par, text):
    """Render **bold**, *italic*, `code` into runs; strip markdown escapes."""
    text = text.replace("\\*", "\u0001")
    for piece in INLINE.split(text):
        if not piece:
            continue
        piece = piece.replace("\u0001", "*")
        if piece.startswith("**") and piece.endswith("**"):
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            par.add_run(piece[1:-1]).italic = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            par.add_run(piece)

def para(text, style=None, align=None, space_after=6, size=None, italic=False):
    p = doc.add_paragraph(style=style)
    runs(p, text)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if size or italic:
        for r in p.runs:
            if size: r.font.size = Pt(size)
            if italic: r.italic = True
    return p

def heading(text, level):
    """A real Heading style, so the section shows up in Word's navigation pane."""
    return doc.add_paragraph(text, style="Heading 1" if level == 1 else "Heading 2")

USABLE_IN = 6.27               # A4 width 8.27" minus 1" margins each side
CHW = 0.059                    # inches per character, Times New Roman 8.5 pt
PAD = 0.16                     # Word's default left+right cell margins
MAXC = 46                      # beyond this a text column just wraps; do not let it eat the table
_BREAKS = re.compile("[ /" + chr(0x2013) + chr(0x2014) + chr(0x2192) + "-]+")

def longest_token(s):
    """The longest run Word cannot break, so a column is never narrower than one of its words."""
    return max([len(x) for x in _BREAKS.split(s) if x] or [1])

def col_widths(rows, ncol):
    """Width every column to fit its longest word, then share out what is left.

    The old rule made width proportional to the longest CELL and then widened column 0 by a
    further quarter. Column 0 holds the row labels, so it took most of the table and the
    numeric columns were left narrower than their own headings: Word broke "position" as
    "positio/n", "interaction" as "interact/ion" and "withheld" as "withhel/d".
    """
    # 8% over the nominal character width: the minus sign, the percent sign and the arrows
    # are wider than a digit, and without the margin Word wrapped "-27.03%" onto two lines.
    need = [max(longest_token(rows[i][j]) for i in range(len(rows))) * CHW * 1.08 + PAD
            for j in range(ncol)]
    want = [min(max(len(rows[i][j]) for i in range(len(rows))), MAXC) * CHW + PAD
            for j in range(ncol)]
    want = [max(w, n) for w, n in zip(want, need)]
    if sum(need) >= USABLE_IN:                 # the minima do not fit; scale them down together
        k = USABLE_IN / sum(need)
        return [Inches(n * k) for n in need]
    if sum(want) <= USABLE_IN:                 # a small table sits at its natural width rather
        return [Inches(w) for w in want]       # than being stretched across the whole page
    slack = USABLE_IN - sum(need)
    room = sum(w - n for w, n in zip(want, need)) or 1.0
    return [Inches(n + slack * (w - n) / room) for n, w in zip(need, want)]

def _border(tag, on):
    e = OxmlElement("w:" + tag)
    e.set(qn("w:val"), "single" if on else "none")
    e.set(qn("w:sz"), "8" if on else "0")
    e.set(qn("w:color"), "000000")
    e.set(qn("w:space"), "0")
    return e

def horizontal_rules(t):
    """Rules above and below the table only. The Guide for Authors asks authors to "Avoid
    vertical rules and shading within table cells"; the Table Grid style drew every one."""
    b = OxmlElement("w:tblBorders")
    for tag, on in (("top", True), ("bottom", True), ("left", False),
                    ("right", False), ("insideH", False), ("insideV", False)):
        b.append(_border(tag, on))
    t._tbl.tblPr.append(b)

def cell_padding(t):
    """A little air above and below each row; the rows were set solid against the rules."""
    m = OxmlElement("w:tblCellMar")
    for tag, w in (("top", "30"), ("bottom", "30"), ("left", "80"), ("right", "80")):
        e = OxmlElement("w:" + tag)
        e.set(qn("w:w"), w); e.set(qn("w:type"), "dxa")
        m.append(e)
    t._tbl.tblPr.append(m)

def md_table(rows):
    """rows: list of list-of-str; row 0 is the header."""
    ncol = max(len(r) for r in rows)
    rows = [r + [""] * (ncol - len(r)) for r in rows]
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = doc.styles["Normal Table"]
    t.autofit = False
    horizontal_rules(t)
    cell_padding(t)
    widths = col_widths(rows, ncol)

    for i, row in enumerate(rows):
        trPr = t.rows[i]._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))      # never break a row across a page
        if i == 0:
            trPr.append(OxmlElement("w:tblHeader"))  # repeat the header on continuation pages
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            c.width = widths[j]
            c.text = ""
            if i == 0:                               # a single rule under the header row
                tb = OxmlElement("w:tcBorders"); tb.append(_border("bottom", True))
                c._tc.get_or_add_tcPr().append(tb)
            p = c.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = (i < len(rows) - 1)
            runs(p, cell)
            for r in p.runs:
                r.font.size = Pt(8.5)
                if i == 0: r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t

def parse(md_text, skip_until=None, stop_at=None):
    """Emit markdown into the document. Returns nothing."""
    lines = md_text.split("\n")
    i, started = 0, skip_until is None
    buf, tbl = [], []

    def flush_par():
        nonlocal buf
        if buf:
            txt = " ".join(buf).strip()
            p = para(txt)
            # A paragraph opening in bold is a panel label; bind it to what follows. The
            # length test this replaces let "(c) Dissociation tested as an interaction — ..."
            # through, and it was left stranded at a page foot with its table overleaf.
            if txt.startswith("**"):
                p.paragraph_format.keep_with_next = True
            buf = []

    def flush_tbl():
        nonlocal tbl
        if tbl:
            md_table(tbl); tbl = []

    while i < len(lines):
        ln = lines[i].rstrip()
        if not started:
            if skip_until and ln.startswith(skip_until):
                started = True
            else:
                i += 1; continue
        if stop_at and ln.startswith(stop_at) and started and not ln.startswith(skip_until or "\0"):
            break
        if ln.startswith("|"):
            flush_par()
            cells = [c.strip() for c in ln.strip("|").split("|")]
            if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
                tbl.append(cells)
            i += 1; continue
        flush_tbl()
        if not ln.strip():
            flush_par(); i += 1; continue
        if ln.startswith("---"):
            flush_par(); i += 1; continue
        if ln.startswith("#"):
            flush_par()
            lvl = len(ln) - len(ln.lstrip("#"))
            txt = ln.lstrip("# ").strip()
            if lvl == 1:
                doc.add_paragraph(txt, style="Title")
            else:
                # each numbered Table starts on a fresh page
                if re.match(r"Table \d+\.", txt):
                    if _seen_table: doc.add_page_break()
                    _seen_table.append(1)
                    heading(txt, 2)          # nested under the "Tables" heading
                else:
                    heading(txt, 1 if lvl == 2 else 2)
            i += 1; continue
        if ln.startswith("> "):
            flush_par()
            p = para(ln[2:], space_after=6, size=11, italic=True)
            p.paragraph_format.left_indent = Inches(0.4)
            i += 1; continue
        buf.append(ln.strip())
        i += 1
    flush_par(); flush_tbl()

# ---------------------------------------------------------------- body
src = io.open(SRC, encoding="utf-8").read()
# drop the internal word-count line
src = re.sub(r"\*\*Word count.*?\n", "", src)

# Assembly order is the Elsevier one -- body, acknowledgements, declarations, references,
# figure captions, tables -- and not the markdown's own order, which puts the figure captions
# before the acknowledgements for the convenience of whoever is editing the text. Parsing the
# file straight through used to strand "Acknowledgements" in the middle of the caption list.
cap_start = src.index("## Figure captions")
ack_start = src.index("## Acknowledgements")
ref_start = src.index("## References")

parse(src[:cap_start])

# ---------------------------------------------------------------- acknowledgements
parse(src[ack_start:ref_start])

# ---------------------------------------------------------------- declarations
# Everything after the reference list, starting at the first heading that follows it. This used
# to slice from "## Conflict of interest", which silently dropped the whole "Data and code
# availability" section -- a declaration the journal requires -- out of the submitted file.
refs = src[ref_start:]
parse(refs[refs.index("\n## ", 1):])
## ", 1):])

# ---------------------------------------------------------------- references
doc.add_page_break()
heading("References", 1)
# The reference list comes from paper_stats so the .docx cannot disagree with the manuscript.
# This loop used to re-split the markdown with a weaker filter, which typeset the horizontal
# rule and the data-availability sentence as two further entries: 24 printed against 22 real.
for block in _references(src):
    p = para(block, space_after=6)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.left_indent = Inches(0.4)

# ---------------------------------------------------------------- figure captions
doc.add_page_break()
heading("Figure captions", 1)
parse(src[cap_start:ack_start], skip_until="**Fig.")

# ---------------------------------------------------------------- tables
doc.add_page_break()
tbl_src = io.open(TBL, encoding="utf-8").read()
tbl_src = tbl_src[tbl_src.index("## Table 1."):]        # drop the file's own title block
heading("Tables", 1)
parse(tbl_src)

# The submission copy carries no figures, as the Guide for Authors requires. Set
# BIOCV_EMBED_FIGS=1 for a reading copy with them appended, for the authors' own checking.
if os.environ.get("BIOCV_EMBED_FIGS"):
    for name, fn in FIGS:
        path = os.path.join(FIGDIR, fn)
        if not os.path.exists(path):
            print("MISSING FIGURE:", path)
            continue
        doc.add_page_break()
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run(name).bold = True
        p1.paragraph_format.space_after = Pt(8)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.line_spacing = 1.0
        p2.add_run().add_picture(path, width=Inches(6.0))
    OUT = OUT.replace(".docx", "_reading_copy.docx")

doc.save(OUT)
print("wrote", OUT, os.path.getsize(OUT), "bytes")

# ---------------------------------------------------------------- figures, separate files
# They used to be pasted at the end of the .docx. The Guide for Authors asks for the opposite:
# artwork "must be supplied as separate files along with the manuscript", named Figure_1 etc.
import shutil
os.makedirs(FIGOUT, exist_ok=True)
for i, (name, fn) in enumerate(FIGS, 1):
    path = os.path.join(FIGDIR, fn)
    if not os.path.exists(path):
        print("MISSING FIGURE:", path); continue
    dst = os.path.join(FIGOUT, f"Figure_{i}.png")
    shutil.copyfile(path, dst)
    print(f"  {name}: {fn} -> {os.path.relpath(dst, 'D:/ROWV_paper')}")
